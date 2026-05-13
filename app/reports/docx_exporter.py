"""
docx_exporter.py
Converts a ReportData object into a richly formatted Word document.

Layout per issue (Bugs / Vulnerabilities / Hotspots):
  🔴 Issue: <message>
  Severity: 🟠 MAJOR
  Occurrences: 171
  Affected Files:
    📄 File: path/to/file.py
    📍 Lines: 63, 75

Ranked Code Smells additionally show:
  🔴 Rank 1: <message>
  ...
"""
from __future__ import annotations

import importlib
import sys
from collections import defaultdict

from app.reports.schemas import ReportData
from app.utils.logging import get_logger

logger = get_logger(__name__)

# ---------------------------------------------------------------------------
# Lazy / safe python-docx import
# ---------------------------------------------------------------------------
_DOCX_AVAILABLE = importlib.util.find_spec("docx") is not None

if _DOCX_AVAILABLE:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
else:
    logger.warning(
        "python-docx is not installed. DOCX generation will be skipped. "
        "Run: pip install python-docx>=1.1.2"
    )

# ---------------------------------------------------------------------------
# Severity → emoji mapping
# ---------------------------------------------------------------------------
EMOJI_MAP: dict[str, str] = {
    "BLOCKER":  "🔴",
    "CRITICAL": "🔴",
    "MAJOR":    "🟠",
    "MINOR":    "🟡",
    "INFO":     "🔵",
    "HIGH":     "🔴",
    "MEDIUM":   "🟠",
    "LOW":      "🟡",
}


def _emoji(severity: str) -> str:
    return EMOJI_MAP.get((severity or "").upper(), "⚪")


# ---------------------------------------------------------------------------
# Grouping helper for flat issue lists (bugs / vulns / hotspots)
# ---------------------------------------------------------------------------
def _group_raw_issues(issue_list: list) -> list:
    """
    Accepts a list of dicts with keys: severity, issue_detail, file_path, line_number.
    Returns grouped list: { message, severity, count, locations: { file_path: [lines] } }
    sorted strongest-severity-first then recurrence-descending.
    """
    WEIGHT = {"BLOCKER": 5, "CRITICAL": 4, "HIGH": 4, "MAJOR": 3, "MEDIUM": 3, "MINOR": 2, "LOW": 2, "INFO": 1}

    groups: dict = defaultdict(lambda: {
        "severity": "INFO",
        "max_weight": 0,
        "count": 0,
        "locations": defaultdict(set)
    })

    for item in issue_list:
        msg = item.get("issue_detail") or "Unknown Issue"
        sev = (item.get("severity") or "INFO").upper()
        weight = WEIGHT.get(sev, 0)
        key = (msg, sev)

        g = groups[key]
        g["count"] += 1
        if weight > g["max_weight"]:
            g["max_weight"] = weight
            g["severity"] = sev

        fp = item.get("file_path") or "Unknown File"
        ln = item.get("line_number") or 0
        g["locations"][fp].add(ln)

    result = []
    for (msg, sev), data in groups.items():
        dict_locs = {fp: sorted(lines) for fp, lines in data["locations"].items()}
        result.append({
            "message": msg,
            "severity": data["severity"],
            "weight": data["max_weight"],
            "count": data["count"],
            "locations": dict_locs,
        })

    result.sort(key=lambda x: (-x["weight"], -x["count"], x["message"]))
    return result


# ---------------------------------------------------------------------------
# Paragraph / formatting helpers
# ---------------------------------------------------------------------------
def _add_heading(doc, text: str, level: int):
    doc.add_heading(text, level=level)


def _add_bold_para(doc, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    return p


def _write_locations(doc, locations: dict[str, list]):
    """Writes the Affected Files block."""
    doc.add_paragraph("Affected Files:")
    for fp, lines in locations.items():
        doc.add_paragraph(f"  📄 File: {fp}")
        visible_lines = [str(ln) for ln in lines if ln and ln > 0]
        if visible_lines:
            doc.add_paragraph(f"  📍 Lines: {', '.join(visible_lines)}")
    doc.add_paragraph()  # breathing room


def _add_issue_section(doc, section_title: str, grouped_issues: list):
    """Renders one full category section (Bugs / Vulns / Hotspots)."""
    _add_heading(doc, section_title, level=1)

    if not grouped_issues:
        doc.add_paragraph("✅ No findings in this category.")
        doc.add_paragraph()
        return

    for item in grouped_issues:
        sev = item["severity"]
        emoji = _emoji(sev)
        _add_bold_para(doc, f"{emoji} Issue: {item['message']}")
        doc.add_paragraph(f"Severity: {emoji} {sev}")
        doc.add_paragraph(f"Occurrences: {item['count']}")
        _write_locations(doc, item["locations"])


# ---------------------------------------------------------------------------
# Main export function
# ---------------------------------------------------------------------------
def _add_comparison_section(doc, diff):
    """Renders the Scan Comparison section. Skipped automatically if diff is None."""
    if diff is None:
        doc.add_paragraph("ℹ️  First scan for this project — no previous scan to compare against.")
        doc.add_paragraph()
        return

    doc.add_paragraph(f"🆕  New Issues           : {len(diff.new_issues)}")
    doc.add_paragraph(f"✅  Resolved Issues       : {len(diff.resolved_issues)}")
    doc.add_paragraph(f"🔁  Persisting Issues     : {diff.persisting_issue_count}")
    doc.add_paragraph(f"📌  Compared Against      : {diff.previous_verification_id}")
    doc.add_paragraph()

    if diff.new_issues:
        _add_heading(doc, "🆕 New Issues (Introduced Since Last Scan)", level=2)
        for issue in diff.new_issues:
            sev = issue.severity or "UNKNOWN"
            emoji = _emoji(sev)
            _add_bold_para(doc, f"{emoji} {issue.type}: {issue.message}")
            doc.add_paragraph(f"Severity: {emoji} {sev}")
            doc.add_paragraph(f"📄 File: {issue.file_path or 'Unknown'}")
            if issue.line_number:
                doc.add_paragraph(f"📍 Line: {issue.line_number}")
            doc.add_paragraph()

    if diff.resolved_issues:
        _add_heading(doc, "✅ Resolved Issues (Fixed Since Last Scan)", level=2)
        for issue in diff.resolved_issues:
            sev = issue.severity or "UNKNOWN"
            emoji = _emoji(sev)
            _add_bold_para(doc, f"{emoji} {issue.type}: {issue.message}")
            doc.add_paragraph(f"📄 File: {issue.file_path or 'Unknown'}")
            doc.add_paragraph()


def export_to_docx(report_data: ReportData, output_path: str) -> None:
    """
    Generate the Word document and save it to output_path.
    Raises RuntimeError if python-docx is unavailable.
    """
    if not _DOCX_AVAILABLE:
        raise RuntimeError(
            "python-docx is not installed. Cannot generate DOCX. "
            "Run: pip install python-docx>=1.1.2"
        )

    doc = Document()

    # ---- Cover / Header ----
    doc.add_heading("Executive Code Quality Report", level=0)
    if report_data.project_name:
        doc.add_paragraph(f"Project         : {report_data.project_name}")
    doc.add_paragraph(f"Verification ID : {report_data.verification_id}")
    doc.add_paragraph()

    # ---- Executive Summary ----
    _add_heading(doc, "Executive Summary", level=1)
    s = report_data.summary
    doc.add_paragraph(f"🐛  Bugs                 : {s.bugs}")
    doc.add_paragraph(f"🔒  Vulnerabilities      : {s.vulnerabilities}")
    doc.add_paragraph(f"🧹  Code Smells          : {s.code_smells}")
    doc.add_paragraph(f"🔥  Security Hotspots    : {s.security_hotspots}")
    doc.add_paragraph()

    # ---- Scan Comparison (new / resolved / persisting) ----
    _add_heading(doc, "Scan Comparison", level=1)
    _add_comparison_section(doc, report_data.diff)
    # ---- Group flat lists globally (by message + severity) ----
    grouped_vulns    = _group_raw_issues(report_data.security_vulnerabilities)
    grouped_bugs     = _group_raw_issues(report_data.bugs_reliability)
    grouped_hotspots = _group_raw_issues(report_data.security_hotspots)

    _add_issue_section(doc, "Security Vulnerabilities", grouped_vulns)
    _add_issue_section(doc, "Reliability Bugs", grouped_bugs)
    _add_issue_section(doc, "Security Hotspots (Manual Review Required)", grouped_hotspots)

    # ---- Top Ranked Code Smells ----
    _add_heading(doc, "Top Repeated Issues (Maintainability)", level=1)

    if not report_data.top_maintainability_issues:
        doc.add_paragraph("✅ No code smells detected.")
    else:
        for item in report_data.top_maintainability_issues:
            sev = item.severity
            emoji = _emoji(sev)
            _add_bold_para(doc, f"{emoji} Rank {item.rank}: {item.issue_detail}")
            doc.add_paragraph(f"Severity: {emoji} {sev}")
            doc.add_paragraph(f"Occurrences: {item.recurrence}")
            _write_locations(doc, item.locations)

    doc.save(output_path)
